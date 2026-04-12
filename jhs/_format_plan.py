# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

base = Path(r'D:/Projects/chengyuan-web/jhs')
dump_path = base / '_plan_final_temp_dump.txt'
out_path = base / '项目计划书_最终版_比赛格式整理.docx'
preview_path = base / '_plan_final_competition_preview.txt'

replacements = {
    '部署��用': '部署应用',
    '产地的���济收益': '产地的经济收益',
    '中���现出色': '中表现出色',
    '完全重��样本': '完全重复样本',
    '与���实值': '与真实值',
    '模��可以': '模型可以',
    '成��分析': '成分分析',
    '两个��段': '两个波段',
    '成分��析': '成分分析',
}

def clean_text(text: str) -> str:
    text = text.replace('　', ' ')
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = text.replace('R平方', 'R²')
    text = text.replace('白利糖度', '°Brix')
    text = text.replace('毫克每100克', 'mg/100g')
    while '  ' in text:
        text = text.replace('  ', ' ')
    return text.strip()

def heading_level(text: str) -> int:
    if text.startswith('第') and '章' in text[:6]:
        return 1
    if text == '参考文献':
        return 1
    head = text.split(' ', 1)[0]
    if not head or not head[0].isdigit():
        return 0
    parts = [p for p in head.split('.') if p]
    if not parts or not all(p.isdigit() for p in parts):
        return 0
    if len(parts) >= 3:
        return 3
    if len(parts) == 2:
        return 2
    return 0

lines = dump_path.read_text(encoding='utf-8').splitlines()
texts = []
for line in lines:
    if ': ' not in line:
        continue
    _, content = line.split(': ', 1)
    texts.append(clean_text(content))

body_start = None
seen_first = False
for idx, text in enumerate(texts):
    if text.startswith('第一章'):
        if not seen_first:
            seen_first = True
        else:
            body_start = idx
            break
if body_start is None:
    body_start = 0
body = [t for t in texts[body_start:] if t and t != '目录']

new_doc = Document()
sec = new_doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.54)
sec.bottom_margin = Cm(2.54)
sec.left_margin = Cm(3.18)
sec.right_margin = Cm(3.18)

styles = new_doc.styles
for style_name in ['Normal', 'Title', 'Heading 1', 'Heading 2', 'Heading 3']:
    style = styles[style_name]
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
styles['Normal'].font.size = Pt(12)
styles['Title'].font.size = Pt(16)
styles['Heading 1'].font.size = Pt(14)
styles['Heading 2'].font.size = Pt(13)
styles['Heading 3'].font.size = Pt(12)


def add_para(text, style=None, center=False, indent=False, bold=False):
    p = new_doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing = 1.5
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if style in (None, 'Normal'):
        run.font.size = Pt(12)
    run.bold = bold or style == 'Title'

add_para('2025年（第18届）中国大学生计算机设计大赛', style='Title', center=True)
add_para('大数据实践赛作品报告', center=True, bold=True)
add_para('作品名称：橙源智鉴——基于多模态光谱融合的海南柑橘品质智能检测与产地溯源系统', center=True)
add_para('整理日期：2026年4月10日', center=True)
new_doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)

inserted_121 = False
for idx, text in enumerate(body):
    if text == '1.2 项目组成' and not inserted_121:
        add_para(text, style='Heading 2')
        add_para('1.2.1 数据层与治理基础', style='Heading 3')
        inserted_121 = True
        continue
    level = heading_level(text)
    if level == 1:
        add_para(text, style='Heading 1')
    elif level == 2:
        add_para(text, style='Heading 2')
    elif level == 3:
        add_para(text, style='Heading 3')
    else:
        add_para(text, indent=True)

new_doc.save(out_path)
with preview_path.open('w', encoding='utf-8') as f:
    for i, p in enumerate(new_doc.paragraphs[:120], 1):
        t = p.text.strip()
        if t:
            f.write(f'{i:03d}: {t}\n')
print(out_path)

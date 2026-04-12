from pathlib import Path
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(r"D:\Projects\chengyuan-web\jhs")
SOURCE = ROOT / "项目计划书_最终版_比赛格式整理.docx"
OUTPUT = ROOT / "项目计划书_最终版_比赛格式整理_按截图要求修正.docx"


def normalize(text: str) -> str:
    return text.replace("\u3000", " ").strip()


def set_text(paragraph, text: str) -> None:
    while paragraph.runs:
        paragraph._p.remove(paragraph.runs[0]._r)
    paragraph.add_run(text)


def insert_paragraph_before(paragraph, text: str, style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph) -> None:
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def replace_with_same_suffix(text: str, old_prefix: str, new_prefix: str) -> str:
    pattern = rf"^{re.escape(old_prefix)}(?P<tail>\s*\d+)?$"
    match = re.match(pattern, normalize(text))
    if not match:
        return text
    return f"{new_prefix}{match.group('tail') or ''}"


doc = Document(str(SOURCE))

# 1) Headings and TOC lines directly mentioned by the screenshots/template.
direct_replacements = {
    "第一章 项目概述": "第1章 作品概述",
    "第一章 作品概述": "第1章 作品概述",
    "第二章 问题分析": "第2章 问题分析",
    "第三章 技术方案": "第3章 技术方案",
    "第四章 系统实现": "第4章 系统实现",
    "第五章 测试分析": "第5章 测试分析",
    "2.1 行业痛点": "2.1 问题来源",
    "2.2 现有方案局限": "2.2 现有解决方案",
    "第六章 项目总结": "第6章 作品总结",
    "第六章 作品总结": "第6章 作品总结",
    "6.1 项目成果": "6.1 作品特色与创新点",
    "6.2 创新点": "6.2 应用推广",
    "6.3 应用价值": "6.3 作品展望",
}

for para in doc.paragraphs:
    text = normalize(para.text)
    if not text:
        continue
    replaced = False
    for old, new in direct_replacements.items():
        new_text = replace_with_same_suffix(text, old, new)
        if new_text != para.text and normalize(new_text) != text:
            set_text(para, new_text)
            replaced = True
            break
    if replaced:
        continue
    if text == "2.3 解决思路":
        set_text(para, "2.4 解决问题的思路")
    elif re.match(r"^2\.3\.\d+\s", text):
        set_text(para, text.replace("2.3.", "2.4.", 1))


paragraphs = doc.paragraphs

# 2) Insert the missing 2.3 section before the old solution section.
solution_heading = None
for para in paragraphs:
    if normalize(para.text) == "2.4 解决问题的思路" and str(para.style.name).startswith("Heading"):
        solution_heading = para
        break

if solution_heading is None:
    raise RuntimeError("未找到 2.4 解决问题的思路 标题，无法继续修正。")

insert_paragraph_before(solution_heading, "2.3 本作品要解决的痛点问题", "Heading 2")
insert_paragraph_before(
    solution_heading,
    "围绕海南柑橘品质检测与产地溯源，当前最突出的矛盾并不是单一算法精度不够，而是样本、指标、光谱和结果之间缺少稳定的一一对应关系。样本登记、滴定指标、液质成分和不同设备的光谱数据往往分散在不同环节，编号口径不一致、采集规范不统一，导致前端展示能讲的内容与后端实验能证明的内容经常脱节。",
    "Normal",
)
insert_paragraph_before(
    solution_heading,
    "第二个痛点在于现有方案多停留在单任务、单模态层面。有人只做糖度预测，有人只做产地分类，也有人只停留在离线实验结果，难以回答“同一批样本还能说明什么”“系统落地后怎么连续使用”这样的问题。对于比赛答辩而言，如果任务之间彼此割裂，评委会很容易追问数据是否共用、结论是否互相支撑。",
    "Normal",
)
insert_paragraph_before(
    solution_heading,
    "第三个痛点是结果表达方式不够面向应用。很多检测工作最终只停留在模型指标表或实验截图，无法直接转成样本说明、批次判断和留档材料。本项目要解决的，正是从统一样本底座、多任务协同分析到结构化结果输出这一整条链路的问题，让每条结论既能在实验中被印证，也能在系统中被看见和被使用。",
    "Normal",
)


# 3) Enrich 2.4 opening so it matches the screenshot's explicit writing requirements.
first_24_subheading = None
for para in doc.paragraphs:
    if normalize(para.text).startswith("2.4.1 "):
        first_24_subheading = para
        break

if first_24_subheading is None:
    raise RuntimeError("未找到 2.4.1 小节，无法补充 2.4 首段。")

insert_paragraph_before(
    first_24_subheading,
    "本项目围绕“同一样本、同一编号、同一数据底座”来组织功能需求与性能目标。功能上，系统需要完成样本录入、光谱关联、任务分析、结果展示、历史回看和报告输出；性能上，要求产地溯源、品质预测和液质成分分析三类任务建立在同一套正式样本划分之上，保证第五章中的测试结果能够相互对应、相互印证。",
    "Normal",
)
insert_paragraph_before(
    first_24_subheading,
    "正式实验数据来自澄迈福橙与琼中绿橙两地产地样本，共形成399个有效样本，其中澄迈199个、琼中200个。数据格式包括4项滴定法指标、36项液质成分数据，以及两套可直接建模的光谱特征：R210光谱228维，S960光谱1402维，拼接后形成1630维双光谱特征。数据获取方式包括产地采样、实验室滴定测定、液质检测和光谱设备采集；在治理过程中，团队排查出HSI模态存在跨产地重复问题，因此正式主表不再将其作为建模依据。",
    "Normal",
)
insert_paragraph_before(
    first_24_subheading,
    "以统一主表中的样本“CM-1”为例，其产地标记为澄迈，对应滴定法指标包括糖度7.35、酸度0.69、糖酸比10.65、维生素C 46.57，同时可关联该样本的R210与S960光谱向量。基于这套统一编号关系，项目进一步开展产地溯源、品质预测和液质成分预测三类任务，避免出现“训练数据一套、展示数据一套”的问题。",
    "Normal",
)
insert_paragraph_before(
    first_24_subheading,
    "在此基础上，项目采用70%、15%、15%的统一划分策略形成训练集、验证集和测试集。后续章节中涉及的准确率、决定系数、平均误差等指标，均建立在这套正式划分与正式主表之上，这也是第五章能够对本章提出的需求与思路进行直接印证的前提。",
    "Normal",
)


def find_heading_index(target: str) -> int:
    for idx, para in enumerate(doc.paragraphs):
        if normalize(para.text) == target and str(para.style.name).startswith("Heading"):
            return idx
    raise RuntimeError(f"未找到标题：{target}")


def replace_section_body(heading_text: str, next_heading_text: str, new_paragraphs: list[str]) -> None:
    start = find_heading_index(heading_text)
    end = find_heading_index(next_heading_text)
    anchor = doc.paragraphs[end]
    current = doc.paragraphs[start + 1 : end]
    for para in reversed(current):
        delete_paragraph(para)
    for text in new_paragraphs:
        insert_paragraph_before(anchor, text, "Normal")


# 4) Rewrite Chapter 6 bodies so the new titles and contents are aligned.
replace_section_body(
    "6.1 作品特色与创新点",
    "6.2 应用推广",
    [
        "本项目最鲜明的特点，是没有把数据整理当成建模前的附属工作，而是先把样本关系理顺，再去谈模型效果。面对多模态数据中出现的异常模态、波段不一致和跨产地重复问题，团队先完成样本核对、模态筛选、波段对齐和统一主表重建，最后才进入正式实验。这样的处理顺序让后续的产地溯源、品质预测和液质成分分析建立在同一套可信样本基础上，结论更容易讲清，也更经得住追问。",
        "第二个特点，是三类任务没有各自为战，而是围绕同一数据底座协同展开。项目把产地识别、理化指标预测和液质成分分析放到同一编号体系下统一推进，使得不同任务之间可以相互补充：产地溯源说明差异确实存在，品质预测给出消费者最关心的口感指标，液质成分进一步解释差异从何而来。相比只做单一分类或单一回归，这种组织方式更接近真实应用场景。",
        "第三个特点，是项目没有停留在实验表格层面，而是把分析过程和结果表达做成了可直接展示的系统形态。网页端和小程序端围绕样本管理、分析发起、结果查看和报告输出构建了完整流程，检测结果不再只是分数和图表，而是能够整理为样本说明、批次记录和正式报告。这样一来，技术结果和应用表达之间不再断层，比赛展示时也更容易形成完整叙事。",
    ],
)

replace_section_body(
    "6.2 应用推广",
    "6.3 作品展望",
    [
        "本项目首先适合落在海南特色柑橘的品质检测与产地说明场景。对于果园、合作社和品牌方来说，样本信息登记、品质初判、结果回看和报告留档都是高频需求，而本项目已经把这些环节串到同一流程中，能够用于批次管理、品牌展示和产地说明。",
        "在分选与流通环节，本项目的价值体现在“判断更快、表达更清楚”。一方面，光谱检测可以减少对纯人工经验的依赖，为产地核验和品质判断提供更稳定的参考；另一方面，系统输出的是结构化结论和规范化页面，不必再把实验结果重新手工整理成说明材料，这对于展示、汇报和留档都更方便。",
        "在教学和科研推广层面，本项目也具备较好的示范意义。它不仅展示了多模态光谱在农产品分析中的应用，还提供了从样本治理、多任务建模到结果表达的完整案例，适合作为大数据实践、智慧农业和农产品质量分析等方向的教学演示与项目延展基础。",
    ],
)

replace_section_body(
    "6.3 作品展望",
    "参考文献",
    [
        "后续迭代将继续围绕两条主线展开。一条主线是扩充样本与品类，在保持现有两地柑橘研究基础的前提下，逐步增加不同年份、不同批次和更多柑橘品种的数据，让模型对产地差异和品质变化的刻画更加充分。另一条主线是继续完善分析维度，把当前已经验证有效的双光谱框架与更多可解释指标结合起来，提升结果表达的丰富度。",
        "系统层面，后续工作将更加重视面向使用场景的细化打磨，包括批次对比、历史追踪、结果检索和模板化输出等功能。这样做的重点不在于把页面做得更复杂，而在于让检测结果更容易被现场人员理解、被管理者采用、被品牌方用于展示。",
        "从更长远的方向看，本项目希望把海南柑橘这条样本链路沉淀为可复用的方法框架：先把数据关系做实，再把多任务结果组织清楚，最后把分析能力转成可传播、可留档、可持续迭代的应用形态。这也是本项目后续继续深化和拓展的核心方向。",
    ],
)


# 5) Add a manual TOC after the正文结构都调整完成.
if not any(normalize(p.text) == "目录" for p in doc.paragraphs[:20]):
    first_heading = None
    for para in doc.paragraphs:
        if normalize(para.text) == "第1章 作品概述" and str(para.style.name) == "Heading 1":
            first_heading = para
            break
    if first_heading is not None:
        toc_lines = [
            ("目录", "Heading 1"),
            ("第1章 作品概述", "Normal"),
            ("1.1 项目背景", "Normal"),
            ("1.2 项目组成", "Normal"),
            ("1.3 创新点", "Normal"),
            ("第2章 问题分析", "Normal"),
            ("2.1 问题来源", "Normal"),
            ("2.2 现有解决方案", "Normal"),
            ("2.3 本作品要解决的痛点问题", "Normal"),
            ("2.4 解决问题的思路", "Normal"),
            ("第3章 技术方案", "Normal"),
            ("3.1 总体技术路线", "Normal"),
            ("3.2 数据治理", "Normal"),
            ("3.3 模型方法", "Normal"),
            ("3.4 多模态融合技术", "Normal"),
            ("第4章 系统实现", "Normal"),
            ("4.1 系统架构", "Normal"),
            ("4.2 核心功能模块", "Normal"),
            ("第5章 测试分析", "Normal"),
            ("5.1 测试设置", "Normal"),
            ("5.2 产地溯源结果", "Normal"),
            ("5.3 品质预测结果", "Normal"),
            ("5.4 液质成分结果", "Normal"),
            ("5.5 综合分析", "Normal"),
            ("第6章 作品总结", "Normal"),
            ("6.1 作品特色与创新点", "Normal"),
            ("6.2 应用推广", "Normal"),
            ("6.3 作品展望", "Normal"),
            ("参考文献", "Normal"),
        ]
        for text, style in toc_lines:
            insert_paragraph_before(first_heading, text, style)


doc.save(str(OUTPUT))
print(str(OUTPUT))

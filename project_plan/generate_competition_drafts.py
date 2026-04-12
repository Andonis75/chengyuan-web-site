from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


BASE_DIR = Path(r"D:\Projects\chengyuan-web\project_plan")

SUMMARY_DOC = BASE_DIR / "computer_design_04_2_summary_draft_20260410.docx"
REPORT_DOC = BASE_DIR / "computer_design_04_3_report_draft_20260410.docx"
NOTES_TXT = BASE_DIR / "computer_design_submission_notes_20260410.txt"


def set_doc_style(doc: Document) -> None:
    styles = doc.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)
    styles["Title"].font.size = Pt(16)
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.size = Pt(12)


def add_paragraph(doc: Document, text: str, style: str | None = None, center: bool = False, indent: bool = False) -> None:
    p = doc.add_paragraph(style=style)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if style == "Title":
        run.bold = True


def build_summary_doc() -> None:
    doc = Document()
    set_doc_style(doc)

    add_paragraph(doc, "中国大学生计算机设计大赛", style="Title", center=True)
    add_paragraph(doc, "作品信息概要表（大数据实践赛草稿）", center=True)
    add_paragraph(doc, "基于现有“橙源智鉴”项目材料整理", center=True)

    fields = [
        ("作品编号", "待确认"),
        ("作品名称", "橙源智鉴——基于多模态光谱融合与深度学习的海南柑橘品质智能检测与产地溯源平台"),
        ("作品大类", "大数据应用"),
        ("作品小类", "实践赛"),
        (
            "作品简介（100字以内）",
            "面向海南橙类品质分级与产地溯源场景，构建集光谱文件上传、智能分析、报告生成、历史归档于一体的平台系统，支持网站端与微信小程序协同使用。",
        ),
        (
            "创新描述（100字以内）",
            "将高光谱分析、等级识别、产地判断与批次档案打通，形成“上传-分析-报告-回看”闭环；采用多端协同与模型服务分层设计，兼顾展示、试点与后续扩展。",
        ),
        (
            "特别说明（100字以内）",
            "当前已完成网站端、小程序端、后端与模型服务原型联动；演示数据下可完成秒级分析流程，部分指标仍需随样本扩充继续验证。",
        ),
    ]

    for label, value in fields:
        add_paragraph(doc, f"{label}：{value}", indent=True)

    add_paragraph(doc, "作者及其分工比例（草稿）", style="Heading 1")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["姓名", "主要分工", "工作量比例"]
    for idx, h in enumerate(headers):
        table.rows[0].cells[idx].text = h
    rows = [
        ("彭斯逸", "项目统筹、算法方案、后端联调、材料总审", "20%"),
        ("唐天赐", "特征提取、指标建模、模型服务开发", "16%"),
        ("艾舜杰", "前端功能集成、小程序交互实现", "14%"),
        ("李明蔚", "全栈架构、可视化页面、结果展示", "14%"),
        ("游亚楠", "市场调研、竞品分析、推广方案", "12%"),
        ("王喆睿", "算法调优、逻辑验证、测试支持", "10%"),
        ("刘东森", "行政协同、进度跟踪、文档整理", "8%"),
        ("饶李充", "财务规划、成本整理、资源对接", "6%"),
    ]
    for name, task, ratio in rows:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = task
        cells[2].text = ratio

    add_paragraph(doc, "指导教师作用（建议勾选项）", style="Heading 1")
    add_paragraph(
        doc,
        "理论指导、技术方案、实验场地、组织协调、其他（商业化与财务指导）。建议对应教师：刘波、龙海侠、程明雄。",
        indent=True,
    )

    add_paragraph(doc, "开发制作平台", style="Heading 1")
    add_paragraph(doc, "Windows、Linux、微信开发者工具。", indent=True)

    add_paragraph(doc, "运行展示平台", style="Heading 1")
    add_paragraph(doc, "Windows、Android、Web浏览器、微信小程序。", indent=True)

    add_paragraph(doc, "开发制作工具", style="Heading 1")
    add_paragraph(
        doc,
        "Next.js 16、React 19、TypeScript、Tailwind CSS 4、ECharts、微信开发者工具、Node.js、Express、Prisma、Python、FastAPI。",
        indent=True,
    )

    add_paragraph(doc, "参考作品（草稿，提交前建议补足具体作品名或文献名）", style="Heading 1")
    refs = [
        "基于高光谱技术的水果无损检测类研究与系统。",
        "农产品质量追溯与产地识别平台类作品。",
        "面向分选与验货场景的智能品控平台类应用。",
    ]
    for idx, ref in enumerate(refs, start=1):
        add_paragraph(doc, f"{idx}. {ref}", indent=True)

    add_paragraph(doc, "提交内容建议", style="Heading 1")
    add_paragraph(
        doc,
        "报告文档、演示视频、PPT、源代码、部署文件、数据集、模型、其他（安装配置说明、用户手册、典型测试样例）。",
        indent=True,
    )

    add_paragraph(doc, "相关文件清单建议", style="Heading 1")
    file_table = doc.add_table(rows=1, cols=4)
    file_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    file_table.style = "Table Grid"
    f_headers = ["序号", "文件与描述", "文件状态", "版权状态"]
    for idx, h in enumerate(f_headers):
        file_table.rows[0].cells[idx].text = h
    file_rows = [
        ("1", "作品信息概要表 PDF（本草稿整理后导出）", "待导出", "自制"),
        ("2", "作品报告 PDF（按04-3模板整理后导出）", "待导出", "自制"),
        ("3", "网站端源代码与部署说明（根目录 Web）", "待整理", "自制"),
        ("4", "小程序源代码与部署说明（xcx/miniprogram）", "待整理", "自制"),
        ("5", "后端服务源代码与接口说明（xcx/server）", "待整理", "自制"),
        ("6", "模型服务样例与接口说明（xcx/model-service）", "待整理", "自制"),
        ("7", "演示视频与答辩PPT", "待补充", "自制"),
        ("8", "典型样本、模型与数据说明（完整版链接待确认）", "待补充", "混合：自制/需标注来源"),
    ]
    for row in file_rows:
        cells = file_table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value

    add_paragraph(doc, "待确认项", style="Heading 1")
    pending = [
        "作品编号。",
        "最终指导教师名单与勾选项。",
        "演示视频、PPT、网盘链接和下载地址。",
        "最终提交的典型数据样本、模型文件范围与版权说明。",
        "参考作品前三项的正式名称。",
    ]
    for item in pending:
        add_paragraph(doc, f"- {item}", indent=True)

    doc.save(SUMMARY_DOC)


def build_report_doc() -> None:
    doc = Document()
    set_doc_style(doc)

    add_paragraph(doc, "2023年（第16届）中国大学生计算机设计大赛", style="Title", center=True)
    add_paragraph(doc, "大数据实践赛作品报告（草稿）", center=True)
    add_paragraph(doc, "作品名称：橙源智鉴——基于多模态光谱融合与深度学习的海南柑橘品质智能检测与产地溯源平台", center=True)
    add_paragraph(doc, "作品编号：待确认", center=True)

    add_paragraph(doc, "第1章 作品概述", style="Heading 1")
    add_paragraph(
        doc,
        "“橙源智鉴”面向海南橙类品质分级、产地识别与结果留档场景，尝试把原本分散的样本采集、光谱分析、人工判断、报告输出和历史回查整合成一套可直接演示和继续试点的数字化系统。项目的创意来源于团队在调研中发现，果园、合作社、分选中心和品牌方都存在一个共同问题：品质判断往往依赖经验，产地说明多停留在纸面，已有检测结果也很难真正进入验货、分级和品牌表达流程。",
        indent=True,
    )
    add_paragraph(
        doc,
        "围绕这一问题，项目构建了网站端、小程序端、后端服务和模型服务组成的多端协同系统。网站端主要承担展示、分析、历史档案和结果可视化，小程序端主要承担采样登记、上传、结果查询和报告回看，后端负责业务流程与数据管理，模型服务负责预测与联调演示。当前系统支持 HDR、RAW、CSV 等文件接入，并能够围绕等级识别、产地判断、品质分析和报告生成形成闭环。",
        indent=True,
    )
    add_paragraph(
        doc,
        "本作品主要面向果园、合作社、分选中心、品牌运营方和区域服务机构等用户群体。对生产端而言，系统的价值在于帮助其形成更稳定的批次记录和分级依据；对品牌端而言，系统的价值在于把样本结果转化为可查询、可留档、可对外说明的材料。项目当前更强调“能展示、能验证、能继续扩样”的阶段性成果，而不是把尚未稳定验证的能力提前写成完全成熟的最终产品。",
        indent=True,
    )

    add_paragraph(doc, "第2章 问题分析", style="Heading 1")
    add_paragraph(doc, "2.1 问题来源", style="Heading 2")
    add_paragraph(
        doc,
        "海南橙类产品在分选、验货和品牌经营过程中，对等级判断、品质说明和产地表达有持续需求，但现实操作中仍较多依赖人工经验与分散记录。随着品牌果、电商果和礼盒果场景增多，结果是否统一、是否能回查、是否能形成批次档案，已经直接影响交易效率和品牌信任。",
        indent=True,
    )
    add_paragraph(doc, "2.2 现有解决方案", style="Heading 2")
    add_paragraph(
        doc,
        "当前常见方案主要有三类。第一类是人工经验分级，优点是上手快，但受人员差异影响大，跨场景复用能力弱。第二类是理化或实验室检测，结果更适合作为验证依据，但周期较长、批量处理能力有限，难以直接承接高频现场判断。第三类是面向外观缺陷或尺寸分选的通用视觉方案，这类系统在外观识别上较有优势，但对内部品质、产地差异和批次档案留存的支撑仍有限。",
        indent=True,
    )
    add_paragraph(
        doc,
        "与上述方案相比，本作品尝试解决的不只是“识别一次结果”，而是把文件接入、分析、报告和历史档案组织成可持续调用的流程，从而让结果真正进入后续经营动作。",
        indent=True,
    )
    add_paragraph(doc, "2.3 本作品要解决的痛点问题", style="Heading 2")
    pain_points = [
        "同一批次样品在不同人员和不同渠道之间容易出现分级口径不一致。",
        "产地说明多依赖纸面材料或口头表达，缺乏可回看、可对照的数字记录。",
        "已有检测结果难以直接进入分选、验货、说明和复盘流程，信息利用率低。",
        "多端协同能力不足，现场采样与后台分析、报告归档之间缺少统一承接。",
    ]
    for item in pain_points:
        add_paragraph(doc, f"- {item}", indent=True)
    add_paragraph(doc, "2.4 解决问题的思路", style="Heading 2")
    add_paragraph(
        doc,
        "本作品采用“光谱文件接入 + 多任务分析 + 报告输出 + 历史留档”的总体思路。数据层面，系统支持 HDR、RAW、CSV 等文件接入；现阶段已沉淀 10 组化验样本、100 个连续波段点位及澄迈福橙、琼中绿橙对比数据，并将继续随试点扩充样本库。模型层面，系统围绕等级识别、产地判断和品质指标估计组织分类与回归任务。应用层面，网站端和小程序端分别服务展示分析与现场操作，最终把结果沉淀为批次记录和标准化报告。",
        indent=True,
    )
    add_paragraph(
        doc,
        "本章提出的“文件接入能力、多任务分析能力、报告输出能力和历史归档能力”，将在第5章通过现有页面联调结果和演示指标进行对应说明。",
        indent=True,
    )

    add_paragraph(doc, "第3章 技术方案", style="Heading 1")
    add_paragraph(
        doc,
        "项目技术方案由网站端、小程序端、业务后端和模型服务四个部分组成。网站端基于 Next.js 16 与 React 19 构建，主要负责首页展示、分析结果可视化、模型说明和历史档案管理；小程序端基于原生微信小程序实现，主要负责采样登记、文件上传、结果查看、报告回看等功能；后端采用 Node.js、Express 与 Prisma 组织业务接口、文件管理和数据存储；模型服务采用 Python 与 FastAPI 形式提供预测接口，用于与业务后端联调。",
        indent=True,
    )
    add_paragraph(
        doc,
        "在数据处理流程上，系统首先完成文件解析与批次关联，再进入预处理、特征组织与模型调用，最后输出等级结论、产地判断、关键指标和文本化报告。为了让结果更适合真实使用场景，项目没有把模型输出停留在单一数值层面，而是继续向报告展示、结果留档和历史回查延伸。",
        indent=True,
    )
    add_paragraph(
        doc,
        "在模型组织上，作品采用分类与回归并行的双任务思路。分类任务用于等级识别和产地判定，回归任务用于糖度、酸度等指标估计；同时保留异常样本提示和可解释展示的接口，便于后续继续扩样和完善判断依据。当前系统中展示的 98.5% 产地识别置信度和 ±0.3% 的 SSC 预测误差，属于现有页面演示口径，后续仍需在更大规模样本下持续验证。",
        indent=True,
    )

    add_paragraph(doc, "第4章 系统实现", style="Heading 1")
    add_paragraph(
        doc,
        "网站端已经实现首页、大屏、分析页、模型页、历史页等核心界面，并将结果卡片、曲线展示、批次对照和历史记录整合到统一风格的交互流程中。用户可在网站端查看分析流程、结果解释、样本档案和报告内容，适合作为项目展示、管理分析和试点对接入口。",
        indent=True,
    )
    add_paragraph(
        doc,
        "小程序端已经实现首页、上传、分析中、结果、报告、历史、样本详情等页面，能够覆盖现场采样与结果回查的主要动作。与单一网页展示不同，小程序端更强调移动场景下的轻量操作，适合在果园、合作社或分选中心进行样本登记和结果查看。",
        indent=True,
    )
    add_paragraph(
        doc,
        "后端服务承担了接口编排、文件处理、批次管理和报告输出等职责；模型服务则提供健康检查与预测接口，形成从业务请求到模型结果返回的完整调用链。当前工程目录中，网站端、后端服务、小程序项目和 Python 模型服务均已独立成形，具备继续联调与部署演示的基础。",
        indent=True,
    )
    add_paragraph(
        doc,
        "在实现过程中，团队重点处理了多端协同、结果口径统一和演示闭环的问题。相比只做单页展示的作品，本项目更强调“页面、数据、接口、报告”之间的一致性，使得同一批次结果可以被现场人员、管理人员和品牌沟通材料共同调用。",
        indent=True,
    )

    add_paragraph(doc, "第5章 测试分析", style="Heading 1")
    add_paragraph(
        doc,
        "现阶段测试主要围绕功能联调、接口可用性和演示指标三类内容展开。功能层面，团队已对网站端首页、分析页、模型页、历史页，以及小程序端上传、分析、结果、报告、历史等主要流程进行联调验证；接口层面，已完成 Node 后端与 Python 模型服务的调用链打通，并可通过健康检查与预测接口验证基本连通性。",
        indent=True,
    )
    add_paragraph(
        doc,
        "在数据与结果层面，当前项目已沉淀 10 组化验样本、100 个连续波段点位及澄迈福橙、琼中绿橙对比数据，并在页面中完成典型演示。根据现有系统展示口径，单次分析流程约为 1.2 秒，产地识别置信度可达 98.5%，SSC 预测误差展示为 ±0.3%。这些结果主要用于说明当前原型系统具备“可联调、可展示、可继续验证”的基础能力，而不是替代更大规模样本验证后的最终性能结论。",
        indent=True,
    )
    add_paragraph(
        doc,
        "从测试结论看，当前系统已经能够证明以下几点：第一，多端协同链路是通的，网站端、小程序端、后端和模型服务能够形成完整业务流程；第二，结果不再停留于单次数值，而是能够进入历史档案与报告输出；第三，项目具备继续扩大样本、补强模型和完善场景试点的工程基础。",
        indent=True,
    )
    add_paragraph(
        doc,
        "与此同时，项目也存在需要继续补强的部分，例如跨批次样本稳定性验证、更多数据规模下的泛化能力评估、不同设备或文件格式的兼容性验证等。后续正式参赛版本应继续补充更完整的实验数据、对比实验与误差分析，以增强报告的说服力。",
        indent=True,
    )

    add_paragraph(doc, "第6章 作品总结", style="Heading 1")
    add_paragraph(doc, "6.1 作品特色与创新点", style="Heading 2")
    features = [
        "围绕真实业务流程组织作品，而不是只展示单一模型结果。",
        "实现网站端、小程序端、后端和模型服务的多端协同。",
        "把分析结果继续延伸到报告输出和历史档案，增强结果复用能力。",
        "在项目表述上强调可验证、可试点、可扩展，避免把阶段性演示结果写成最终定论。",
    ]
    for item in features:
        add_paragraph(doc, f"- {item}", indent=True)
    add_paragraph(doc, "6.2 应用推广", style="Heading 2")
    add_paragraph(
        doc,
        "本作品适合首先进入果园、合作社、分选中心和品牌方等场景，作为样本登记、批次分析、结果说明和档案留存工具使用。若后续样本规模和验证深度继续提高，系统还可向区域品牌服务机构、科研团队和相关监管辅助场景扩展。",
        indent=True,
    )
    add_paragraph(doc, "6.3 作品展望", style="Heading 2")
    add_paragraph(
        doc,
        "后续工作将重点放在四个方向：继续补强区域样本库与理化标签；提升跨批次、跨场景的模型稳定性；完善报告模板与可解释分析；进一步梳理数据、模型、部署文件和版权说明，使作品更符合正式参赛提交要求。",
        indent=True,
    )

    add_paragraph(doc, "参考文献（草稿，提交前建议按最终采用文献统一格式校核）", style="Heading 1")
    references = [
        "农业农村部. 全国智慧农业行动计划（2024—2028年）[Z]. 2024.",
        "中华人民共和国农产品质量安全法[Z]. 2022年修订.",
        "国家数据局等17部门. “数据要素×”三年行动计划（2024—2026年）[Z]. 2024.",
        "海南省统计局. 海南省2024年国民经济和社会发展统计公报[Z]. 2025.",
        "高光谱成像与农产品无损检测相关研究文献[待补充最终引用].",
    ]
    for idx, ref in enumerate(references, start=1):
        add_paragraph(doc, f"[{idx}] {ref}", indent=True)

    doc.save(REPORT_DOC)


def write_notes() -> None:
    notes = """计算机设计大赛材料整理说明（2026-04-10）

已生成：
1. 04-2 概要表草稿
2. 04-3 作品报告草稿

本轮已按模板落实的内容：
- 作品名称与项目定位
- 作品大类/小类建议（大数据应用 / 实践赛）
- 基于当前项目的功能、技术栈、系统结构和测试口径
- 与现有项目材料一致的团队分工草稿

仍需最终确认的内容：
- 作品编号
- 最终作者名单与分工比例是否调整
- 指导教师最终名单及勾选项
- 演示视频、答辩PPT、网盘地址、下载地址
- 数据集/模型的最终提交范围与版权说明
- 参考作品前三项的正式名称
- 参考文献中的技术类文献补全

建议下一步：
1. 先审 04-2，确认名称、分工、提交清单。
2. 再审 04-3，优先补强第5章测试分析中的数据支撑。
3. 最后统一导出 PDF 并按 04-1 要求整理目录结构。
"""
    NOTES_TXT.write_text(notes, encoding="utf-8")


def main() -> None:
    build_summary_doc()
    build_report_doc()
    write_notes()
    print(SUMMARY_DOC)
    print(REPORT_DOC)
    print(NOTES_TXT)


if __name__ == "__main__":
    main()

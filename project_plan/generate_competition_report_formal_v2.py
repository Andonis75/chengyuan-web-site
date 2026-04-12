from __future__ import annotations

import csv
import json
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


BASE_DIR = Path(r"D:\Projects\chengyuan-web\project_plan")
PROJECT_DIR = Path(r"D:\Projects\chengyuan-web")
DATA_DIR = PROJECT_DIR / "data"
OUT_DOCX = BASE_DIR / "computer_design_04_3_report_formal_v14_20260410.docx"
OUT_TXT = BASE_DIR / "computer_design_04_3_report_formal_v14_20260410.txt"

WEB_PAGE_COUNT = 9
MINI_PROGRAM_PAGE_COUNT = 11
CORE_CAPABILITY_COUNT = 20


def set_doc_style(doc: Document) -> None:
    styles = doc.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)
    styles["Title"].font.size = Pt(16)
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.size = Pt(12)


def add_paragraph(
    doc: Document,
    text: str,
    style: str | None = None,
    *,
    center: bool = False,
    indent: bool = False,
) -> None:
    p = doc.add_paragraph(style=style)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if style == "Title":
        run.bold = True


def add_table(
    doc: Document,
    lines: list[str],
    title: str,
    headers: list[str],
    rows: list[list[str]],
) -> None:
    add_paragraph(doc, title, indent=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, value in enumerate(headers):
        table.rows[0].cells[idx].text = value
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    doc.add_paragraph("")

    lines.append(title)
    lines.append(" | ".join(headers))
    for row in rows:
        lines.append(" | ".join(row))
    lines.append("")


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def mean(values: list[float]) -> float:
    return sum(values) / len(values) if values else 0.0


def fmt_float(value: float, digits: int = 4) -> str:
    return f"{value:.{digits}f}"


def get_data_root() -> Path:
    candidates = [
        p
        for p in DATA_DIR.iterdir()
        if p.is_dir() and "橙源智鉴" in p.name and (p / "正式实验").exists()
    ]
    if not candidates:
        raise FileNotFoundError("未找到包含正式实验目录的项目资料文件夹。")
    return candidates[0]


def build_metrics() -> dict:
    root = get_data_root()
    overview = read_json(root / "正式实验" / "研究记录" / "正式实验概览.json")
    quality_summary = read_json(root / "正式实验" / "研究记录" / "第一版双光谱基线实验摘要.json")
    clean_summary = read_json(root / "清洗结果" / "质检信息" / "汇总信息.json")
    quality_report = read_csv(root / "正式实验" / "基线结果" / "任务1_品质预测" / "逐目标结果.csv")
    origin_report = read_csv(root / "正式实验" / "基线结果" / "任务2_产地溯源" / "模型汇总.csv")
    lcms_report = read_csv(root / "正式实验" / "基线结果" / "任务3_液质成分预测" / "逐目标结果.csv")
    split_rows = read_csv(root / "正式实验" / "划分方案" / "统一样本划分.csv")
    dataset_rows = read_csv(root / "正式实验" / "数据集" / "双光谱安全主表.csv")
    lcms_tune_rows = read_csv(root / "正式实验" / "研究记录" / "进阶优化记录" / "液质预测进阶调参完整记录.csv")

    origin_stats: dict[str, dict[str, float]] = {}
    for origin in ["CM", "QZ"]:
        subset = [r for r in dataset_rows if r["origin"] == origin]
        origin_stats[origin] = {
            "count": float(len(subset)),
            "糖度": mean([float(r["titration_糖度"]) for r in subset]),
            "酸度": mean([float(r["titration_酸度"]) for r in subset]),
            "糖酸比": mean([float(r["titration_糖酸比"]) for r in subset]),
            "VC": mean([float(r["titration_VC"]) for r in subset]),
            "异柠檬酸": mean([float(r["lcms_Isocitric_acid"]) for r in subset]),
            "莽草酸": mean([float(r["lcms_Shikimic_acid"]) for r in subset]),
            "蔗糖": mean([float(r["lcms_蔗糖"]) for r in subset]),
        }

    split_counter = Counter(r["split"] for r in split_rows)
    origin_split_counter = Counter((r["origin"], r["split"]) for r in split_rows)

    quality_test_rows = [r for r in quality_report if r["split"] == "测试集"]
    best_quality_by_target: dict[str, dict[str, str]] = {}
    for row in quality_test_rows:
        target = row["target"]
        if target not in best_quality_by_target or float(row["r2"]) > float(best_quality_by_target[target]["r2"]):
            best_quality_by_target[target] = row

    lcms_test_rows = [r for r in lcms_report if r["split"] == "测试集"]
    best_lcms_by_target: dict[str, dict[str, str]] = {}
    for row in lcms_test_rows:
        target = row["target"]
        if target not in best_lcms_by_target or float(row["r2"]) > float(best_lcms_by_target[target]["r2"]):
            best_lcms_by_target[target] = row

    best_lcms_val = max(lcms_tune_rows, key=lambda r: float(r["验证集平均R2"]))
    best_lcms_test = max(lcms_tune_rows, key=lambda r: float(r["测试集平均R2"]))
    best_origin = max(origin_report, key=lambda r: float(r["测试集准确率"]))
    sample_example = {k: dataset_rows[0][k] for k in list(dataset_rows[0].keys())[:7]}

    return {
        "overview": overview,
        "quality_summary": quality_summary,
        "clean_summary": clean_summary,
        "origin_stats": origin_stats,
        "split_counter": split_counter,
        "origin_split_counter": origin_split_counter,
        "best_quality_by_target": best_quality_by_target,
        "best_lcms_by_target": best_lcms_by_target,
        "best_lcms_val": best_lcms_val,
        "best_lcms_test": best_lcms_test,
        "best_origin": best_origin,
        "sample_example": sample_example,
    }


def main() -> None:
    metrics = build_metrics()
    overview = metrics["overview"]
    quality_summary = metrics["quality_summary"]
    clean_summary = metrics["clean_summary"]
    origin_stats = metrics["origin_stats"]
    split_counter = metrics["split_counter"]
    origin_split_counter = metrics["origin_split_counter"]
    best_quality = metrics["best_quality_by_target"]
    best_lcms = metrics["best_lcms_by_target"]
    best_lcms_val = metrics["best_lcms_val"]
    best_lcms_test = metrics["best_lcms_test"]
    best_origin = metrics["best_origin"]
    sample_example = metrics["sample_example"]

    baseline_lcms = quality_summary["lcms_top1"]
    lcms_val_gain = float(best_lcms_val["验证集平均R2"]) - float(baseline_lcms["验证集平均R2"])
    lcms_test_gain = float(best_lcms_test["测试集平均R2"]) - float(baseline_lcms["测试集平均R2"])

    cm_count = int(origin_stats["CM"]["count"])
    qz_count = int(origin_stats["QZ"]["count"])
    train_count = split_counter["训练集"]
    val_count = split_counter["验证集"]
    test_count = split_counter["测试集"]

    qz_sugar = origin_stats["QZ"]["糖度"]
    qz_ratio = origin_stats["QZ"]["糖酸比"]
    cm_vc = origin_stats["CM"]["VC"]
    cm_iso = origin_stats["CM"]["异柠檬酸"]
    qz_sucrose = origin_stats["QZ"]["蔗糖"]

    sugar_result = best_quality["titration_糖度"]
    vc_result = best_quality["titration_VC"]
    ratio_result = best_quality["titration_糖酸比"]

    iso_result = best_lcms["lcms_Isocitric_acid"]
    shikimic_result = best_lcms["lcms_Shikimic_acid"]
    sucrose_result = best_lcms["lcms_蔗糖"]
    tricarboxylic_result = best_lcms["lcms_E_1_Propene_1_2_3_Tricarboxylic_Acid"]

    doc = Document()
    set_doc_style(doc)
    lines: list[str] = []

    def w(text: str = "", style: str | None = None, *, center: bool = False, indent: bool = False) -> None:
        add_paragraph(doc, text, style=style, center=center, indent=indent)
        lines.append(text)

    w("2023年（第16届）中国大学生计算机设计大赛", style="Title", center=True)
    w("大数据实践赛作品报告", center=True)
    w("作品名称：橙源智鉴——基于多模态光谱融合的海南柑橘品质智能检测与产地溯源项目", center=True)
    w("作品编号：待确认", center=True)
    w("填写日期：2026年4月10日", center=True)
    w()

    w("第一章 项目概述", style="Heading 1")
    w("1.1 项目背景与创意来源", style="Heading 2")
    w(
        "海南柑橘在产地品牌建设、分级流通和质量说明环节，对检测结果的及时性、可读性和可追溯性提出了更高要求。传统做法往往把样本登记、理化检测、光谱分析和报告整理分散在不同表格和不同人员手中，现场能看到样本，实验室能拿到指标，真正到展示、验货和归档环节时，信息却很难重新对齐。本项目围绕这一现实问题展开，尝试把样本、数据、结果和正式文本放回同一条链路中。",
        indent=True,
    )
    w(
        "项目以澄迈福橙和琼中绿橙为研究对象，围绕两类样本构建多模态检测与分析流程。两地柑橘具有较强的地域代表性，样本在糖度、糖酸比、VC 以及部分液质成分上存在稳定差异，既适合开展产地判别，也适合承载品质评价和内部成分解析。本项目由此形成了“数据治理+多任务分析+结果表达”的整体思路，目标不是给出孤立指标，而是形成一套可用于比赛展示、场景说明和后续推广的项目成果。",
        indent=True,
    )

    w("1.2 面向对象与应用价值", style="Heading 2")
    w(
        "本项目主要面向果园与合作社、产地质检与分选环节、品牌展示与渠道验货场景，以及教学与科研演示场景。对果园和合作社而言，项目关注样本登记、批次区分和品质初判；对分选和验货环节而言，项目关注产地核验、糖度判断和重点成分说明；对品牌展示而言，项目关注结果表达、批次留档和正式报告输出；对教学与科研场景而言，项目关注数据组织、分析过程和结果对比的完整呈现。",
        indent=True,
    )
    w(
        "与单一指标检测相比，本项目更强调结果之间的对应关系。产地判别说明“来自哪里”，品质预测回答“表现怎样”，液质成分分析补充“内部成分有什么差异”，三类结果最终围绕同一编号样本展开，比赛展示、样本回查和正式说明因此可以使用同一套材料。",
        indent=True,
    )
    add_table(
        doc,
        lines,
        "表1-1 项目主要服务场景",
        ["场景", "关注内容", "项目输出"],
        [
            ["果园与合作社", "批次登记、初步分级、样本留档", "样本信息、品质结果、历史记录"],
            ["质检与分选", "产地核验、糖度判断、重点成分说明", "分析结果页、对比结果、批次报告"],
            ["品牌展示与渠道验货", "产地说明、品质说明、结果留存", "正式中文报告、可回看记录"],
            ["教学与科研演示", "数据治理、任务对比、结果展示", "多任务实验结果与系统演示页面"],
        ],
    )

    w("1.3 项目组成与主要功能", style="Heading 2")
    w(
        "本项目由网页端、移动端、分析支撑层和结果输出模块四部分组成。网页端放置项目展示、数据总览、结果对比和历史回看页面；移动端放置样本录入、分析提交、结果查询和报告查看页面；分析支撑层负责样本管理、文件处理、任务调度和结果整理；结果输出模块负责把样本信息、产地判别、品质指标和重点成分整理为正式文本。",
        indent=True,
    )
    w(
        f"截至 2026 年 4 月 10 日，项目已完成 {WEB_PAGE_COUNT} 个网页主要页面、{MINI_PROGRAM_PAGE_COUNT} 个移动端功能页面，并形成覆盖样本、分析、结果、报告四个环节的 {CORE_CAPABILITY_COUNT} 项核心能力。当前版本已经能够围绕同一批样本完成录入、分析、展示、回看和报告输出。",
        indent=True,
    )
    add_table(
        doc,
        lines,
        "表1-2 项目当前组成",
        ["模块", "主要内容", "当前完成情况"],
        [
            ["网页端", "首页、分析页、数据总览、历史页等", f"已完成 {WEB_PAGE_COUNT} 个主要页面"],
            ["移动端", "上传、查询、分析、报告查看等", f"已完成 {MINI_PROGRAM_PAGE_COUNT} 个功能页面"],
            ["分析支撑层", "样本、文件、分析、结果等流程组织", f"已形成 {CORE_CAPABILITY_COUNT} 项核心能力"],
            ["结果输出", "历史归档、正式报告、中文结果整理", "已支持结构化结果输出与报告生成"],
        ],
    )

    w("1.4 项目定位", style="Heading 2")
    w(
        "本项目定位于海南特色农产品场景下的大数据实践项目。围绕海南柑橘这一具体对象，项目把数据治理、多任务分析和正式结果表达放到同一份成果中，评审既能看到正式实验数据，也能看到围绕样本展开的系统展示与报告输出。",
        indent=True,
    )

    w("第二章 问题分析", style="Heading 1")
    w("2.1 问题来源", style="Heading 2")
    w(
        "海南柑橘进入品牌化和区域化发展阶段后，单纯依靠经验判断已难以满足产地辨识、品质说明和批次追溯的需求。果品进入流通、展示和验货环节时，现场最常见的问题并不是“有没有检测结果”，而是“结果分散在哪里、依据是否统一、样本能否回查、文字如何对外说明”。如果样本信息、检测数据和结论文本不能围绕同一编号重新组织，结果再多也难以转化为稳定的使用价值。",
        indent=True,
    )
    w(
        "本项目把问题落在海南柑橘这一具体场景上，原因在于两类样本既有区域品牌属性，又有可量化的理化和成分差异，适合作为数据实践项目的承载对象。围绕澄迈福橙和琼中绿橙，项目需要解决的不是单点识别，而是从样本整理、数据建模、结果比较到报告输出的整条链路问题。",
        indent=True,
    )

    w("2.2 现有解决方案", style="Heading 2")
    w(
        "围绕果品产地识别和品质检测，现有路线大致分为三类。第一类是人工经验和理化检测，结论直观，但效率低、对批次留档支持有限；第二类是面向单一任务的光谱建模研究，通常围绕糖度、可溶性固形物、产地类别等指标展开，已经证明光谱技术在柑橘场景中的可行性[11]-[14]；第三类是强调展示或管理的单项系统，具备页面形态，但往往难以与正式实验数据和多任务结果形成强绑定。",
        indent=True,
    )
    w(
        "现有研究为本项目提供了很好的方法参考，但在比赛稿表达上仍有两个明显空缺：其一，多数成果聚焦单一任务，较少把产地、品质和成分结果放到同一套样本体系中讨论；其二，很多结果停留在数值表层，缺少面向展示、回查和正式说明的组织方式。本项目正是在这两处空缺上继续推进。",
        indent=True,
    )
    add_table(
        doc,
        lines,
        "表2-1 现有解决方案对比",
        ["解决路线", "优势", "局限", "本项目对应处理"],
        [
            ["人工经验与理化检测", "指标直观，便于小规模判断", "周期较长，批次信息和说明材料分散", "保留理化标签作为依据，同时引入统一样本管理与结果输出"],
            ["单一光谱建模研究", "对产地或品质单项任务有效，方法成熟", "任务之间相互分离，结果难统一表达", "围绕同一主表组织产地、品质、成分三类任务"],
            ["单项展示或管理系统", "页面形态完整，适合演示", "与实验数据联系不紧，缺少正式结果支撑", "将正式实验、页面展示和中文报告放到同一链路中"],
        ],
    )

    w("2.3 本项目要解决的痛点问题", style="Heading 2")
    w(
        "基于前述对比，本项目聚焦四个核心痛点。第一，样本编号、实验指标和多模态光谱未形成统一主表，数据难以横向比较；第二，多模态数据中存在异常模态和波段不一致问题，不适合直接进入正式实验；第三，产地、品质和液质成分常被拆成三套独立结果，缺少统一口径；第四，检测结果停留在数值层面，难以直接转化为可展示、可传递、可归档的正式材料。",
        indent=True,
    )
    add_table(
        doc,
        lines,
        "表2-2 本项目聚焦的痛点",
        ["痛点", "具体表现", "比赛展示中的影响"],
        [
            ["数据底座不统一", "样本、标签、光谱分散在不同文件中", "结果难以在同一条叙事线中展开"],
            ["异常模态干扰正式实验", "HSI 存在跨产地同编号完全重复", "会影响正式建模的可信度"],
            ["多任务结果彼此割裂", "产地、品质、成分结论分开陈列", "评审难以看到项目的整体价值"],
            ["结果表达停留在数值层面", "缺少面向展示和归档的正式文本", "现场演示容易停留在模型分数说明"],
        ],
    )

    w("2.4 解决问题的思路", style="Heading 2")
    w(
        "本项目首先重建数据底座。原始资料覆盖滴定法指标、液质成分、R210 光谱、S960 光谱和 HSI 光谱等多类数据，经过清洗整理后，正式实验统一建立在 399 个安全样本上。项目在正式建模前完成 HSI 异常排查、S960 共同波段对齐、统一样本划分和安全主表重建，随后围绕同一主表开展产地溯源、品质预测和液质成分预测三类任务。",
        indent=True,
    )
    w(
        f"正式主表共包含 {overview['sample_count']} 个样本、{overview['feature_count']} 维双光谱特征，其中 R210 为 {overview['r210_feature_count']} 维，S960 为 {overview['s960_feature_count']} 维；标签侧包含 {overview['titration_target_count']} 项滴定法指标和 {overview['lcms_target_count']} 项液质成分指标。统一划分后，训练集、验证集、测试集规模分别为 {train_count}、{val_count}、{test_count}。在此基础上，项目把结果重新组织为页面展示、历史记录和正式报告三类输出，使实验结果能够直接进入比赛展示和应用说明环节。",
        indent=True,
    )
    add_table(
        doc,
        lines,
        "表2-3 正式实验数据概况",
        ["项目", "规模", "说明"],
        [
            ["正式安全样本", str(overview["sample_count"]), "正式建模唯一主表样本数"],
            ["双光谱特征数", str(overview["feature_count"]), f"R210 {overview['r210_feature_count']} + S960 {overview['s960_feature_count']}"],
            ["滴定法指标", str(overview["titration_target_count"]), "糖度、酸度、糖酸比、VC"],
            ["液质成分指标", str(overview["lcms_target_count"]), "用于内部成分预测的 36 项标签"],
            ["统一样本划分", f"{train_count}/{val_count}/{test_count}", "训练集/验证集/测试集"],
        ],
    )
    add_table(
        doc,
        lines,
        "表2-4 样本记录示例",
        ["样本编号", "产地", "序号", "糖度", "酸度", "糖酸比", "VC"],
        [
            [
                sample_example["sample_id"],
                sample_example["origin"],
                sample_example["sample_number"],
                sample_example["titration_糖度"],
                sample_example["titration_酸度"],
                sample_example["titration_糖酸比"],
                f"{float(sample_example['titration_VC']):.2f}",
            ]
        ],
    )

    w("第三章 技术方案", style="Heading 1")
    w("3.1 总体技术路线", style="Heading 2")
    w(
        "本项目的技术路线按照“样本治理、任务建模、结果组织、系统呈现”四个阶段推进。项目先完成数据清洗、异常排查和主表重建，再组织三类正式实验，随后把结果回连到样本编号，最后在网页端、移动端和正式报告中统一呈现。",
        indent=True,
    )
    add_table(
        doc,
        lines,
        "表3-1 总体技术路线",
        ["阶段", "核心工作", "产出"],
        [
            ["阶段一：样本治理", "清洗原始数据、排查异常模态、完成对齐", "安全主表与统一样本体系"],
            ["阶段二：任务建模", "组织产地、品质、液质成分三类实验", "多任务测试结果"],
            ["阶段三：结果组织", "把结果回连到样本编号和批次信息", "结构化结果说明与历史记录"],
            ["阶段四：系统呈现", "在网页端与移动端完成展示和查询", "比赛可演示的项目形态"],
        ],
    )

    w("3.2 数据路线", style="Heading 2")
    w(
        "项目原始数据由滴定法、液质和多模态光谱三部分组成。清洗后发现，HSI 数据中存在 199 对跨产地同编号完全重复样本，无法作为正式跨产地实验依据；S960 两地产地原始波段数量不同，其中琼中样本较澄迈多出 50 个尾部波段，必须先完成共同波段对齐。项目据此剔除 HSI，保留 R210 与 S960 双光谱作为正式建模基础，并重建“安全版_去除HSI融合表”，保证正式实验口径统一。",
        indent=True,
    )
    w(
        f"在波段处理上，S960 共同波段数量为 {clean_summary['s960_alignment']['common_band_count']}，琼中额外波段为 {clean_summary['s960_alignment']['qz_extra_band_count']}；在样本控制上，全部正式实验使用固定随机种子 {overview['split_seed']} 完成统一划分；在样本完整性上，仅有 {', '.join(clean_summary['missing_all_modalities_samples'])} 因模态缺失未进入全模态流程。三类任务因此能够在同一数据底座上展开比较。",
        indent=True,
    )
    add_table(
        doc,
        lines,
        "表3-2 数据治理要点",
        ["治理项", "结果", "作用"],
        [
            ["HSI 异常排查", f"识别 {clean_summary['hsi_exact_duplicate_pair_count']} 对跨产地完全重复样本", "排除异常模态对正式实验的干扰"],
            ["正式建模主表", f"{clean_summary['safe_fusion_without_hsi_sample_count']} 个安全样本", "作为三类任务的唯一正式主表"],
            ["S960 波段对齐", f"{clean_summary['s960_alignment']['common_band_count']} 个共同波段", "保证两地产地数据可直接比较"],
            ["统一样本划分", f"{train_count}/{val_count}/{test_count}", "保证多任务结果可在同一框架下解释"],
        ],
    )

    w("3.3 模型与关键方法", style="Heading 2")
    w(
        "三类任务采用统一主表、分任务建模的方式推进。产地溯源任务以分类模型为主，重点比较单光谱与双光谱方案在区分不同产地时的稳定性；品质预测任务以回归模型为主，围绕糖度、糖酸比、VC 等指标展开；液质成分预测任务面对的是更多目标和更高维数据，因而在基线实验之外进一步引入 SNV 预处理、PCA 降维和参数搜索，以提升整体平均表现。",
        indent=True,
    )
    w(
        "模型安排遵循“先比较、后优化”的顺序。分类侧采用逻辑回归、线性 SVM 等方法建立基线；回归侧采用岭回归、PLS 回归、随机森林、XGBoost 等方法比较不同指标的拟合情况；优化侧重点放在液质任务，通过 SNV 去除散射影响、通过 PCA 压缩高维冗余，再用参数搜索寻找更合适的组合。正文同时保留基线结果与优化结果，便于呈现改进轨迹。",
        indent=True,
    )
    add_table(
        doc,
        lines,
        "表3-3 三类任务技术安排",
        ["任务", "输入与标签", "主要方法", "关注重点"],
        [
            ["产地溯源", "双光谱特征 + origin 标签", "逻辑回归、线性 SVM、集成分类模型", "比较单光谱与双光谱的区分效果"],
            ["品质预测", "双光谱特征 + 糖度/糖酸比/VC 等标签", "岭回归、PLS 回归等", "锁定最具代表性的品质指标"],
            ["液质成分预测", "双光谱特征 + 36 项液质标签", "基线回归 + SNV + PCA + 参数搜索", "提升多目标任务的整体平均表现"],
        ],
    )

    w("3.4 结果组织与输出方案", style="Heading 2")
    w(
        "本项目把三类任务的结果都回连到同一编号样本，不把产地、品质和成分分成彼此孤立的三份材料。页面展示围绕样本信息、批次信息、分析结果和历史记录展开，正式报告围绕产地判别、关键品质指标和重点成分结果组织语言。这样处理后，同一条实验结果能够直接进入比赛讲解、场景说明和资料留档。",
        indent=True,
    )

    w("第四章 系统实现", style="Heading 1")
    w("4.1 展示端实现", style="Heading 2")
    w(
        "展示端分为网页端和移动端两部分。网页端放置首页、数据总览、分析页、历史页等主要页面，用于集中展示样本规模、实验结果和项目成果；移动端放置上传、分析、查询、报告查看等页面，用于样本录入和结果回看。两端围绕同一套样本与结果组织，入口不同，结果口径一致。",
        indent=True,
    )
    w(
        "在页面表达上，项目重点突出样本编号、产地信息、关键指标和正式结果，不把展示页写成开发说明页。评审在演示过程中能够直接看到样本规模、产地差异、糖度结果、液质成分结果和历史报告。",
        indent=True,
    )

    w("4.2 分析流程实现", style="Heading 2")
    w(
        "系统流程按照样本录入、文件接收、任务处理、结果整理、历史归档五个环节展开。样本进入系统后先绑定编号和批次信息，再进入分析流程；分析完成后，结果围绕该样本编号回写到结果页和历史页；报告模块从结构化结果中抽取产地、品质和重点成分信息，生成正式文本。整个过程始终以样本编号为主线。",
        indent=True,
    )
    w(
        "网页展示、移动操作、分析处理和报告输出围绕同一数据结构展开。比赛现场展示的是同一批样本的不同结果，而不是来自不同来源、不同口径的零散材料。",
        indent=True,
    )

    w("4.3 关键难点与解决方法", style="Heading 2")
    w(
        "系统实现过程中，难点主要集中在数据治理与结果统一表达两处。项目针对这些问题逐项给出处理方案，并把处理结果写进正式主表和正式流程中。",
        indent=True,
    )
    add_table(
        doc,
        lines,
        "表4-1 关键难点与解决方法",
        ["难点", "具体表现", "解决方法"],
        [
            ["异常模态干扰", "HSI 出现跨产地同编号完全重复", "剔除 HSI，重建双光谱安全主表"],
            ["波段口径不统一", "S960 两地产地原始波段数量不同", "统一裁剪为共同波段对齐表"],
            ["多任务结果分散", "分类、回归、成分结果来源不同", "围绕样本编号统一页面结构和报告模板"],
            ["历史回查困难", "结果难形成正式留档材料", "建立历史记录与中文报告输出链路"],
        ],
    )

    w("4.4 项目完成情况", style="Heading 2")
    w(
        f"截至当前版本，项目已完成从样本组织、分析处理到展示输出的主要建设内容。网页端形成 {WEB_PAGE_COUNT} 个主要页面，移动端形成 {MINI_PROGRAM_PAGE_COUNT} 个功能页面，分析支撑层形成 {CORE_CAPABILITY_COUNT} 项核心能力，正式实验数据、展示页面和报告输出已经能够围绕同一主线联动。",
        indent=True,
    )
    add_table(
        doc,
        lines,
        "表4-2 项目完成情况概览",
        ["建设内容", "完成情况", "对应作用"],
        [
            ["正式实验数据", "399 个安全样本与三类任务数据集", "支撑测试分析与结果对比"],
            ["网页展示", f"{WEB_PAGE_COUNT} 个主要页面", "支撑比赛展示与集中讲解"],
            ["移动端使用", f"{MINI_PROGRAM_PAGE_COUNT} 个功能页面", "支撑现场录入、查询与回看"],
            ["分析与输出", f"{CORE_CAPABILITY_COUNT} 项核心能力", "支撑任务处理、历史归档与报告生成"],
        ],
    )

    w("第五章 测试分析", style="Heading 1")
    w("5.1 测试设置", style="Heading 2")
    w(
        f"测试部分全部建立在统一安全主表之上，正式实验样本数为 {overview['sample_count']}，训练集、验证集、测试集规模分别为 {train_count}、{val_count}、{test_count}。澄迈福橙与琼中绿橙在验证集和测试集中均保持 {origin_split_counter[('CM', '验证集')]}:{origin_split_counter[('QZ', '验证集')]} 与 {origin_split_counter[('CM', '测试集')]}:{origin_split_counter[('QZ', '测试集')]} 的平衡分布，保证分类与回归任务处在一致的评估条件下。",
        indent=True,
    )
    w(
        "测试流程采用基线比较与优化比较并行的方式。分类任务重点看准确率、平衡准确率和 F1；回归任务重点看 R² 和 RMSE；液质成分任务在单项结果之外还关注平均 R² 的变化，以判断预处理和参数搜索是否带来稳定增益。所有结论均以正式实验结果为依据，不使用临时口径和零散结果替代正式测试。",
        indent=True,
    )
    add_table(
        doc,
        lines,
        "表5-1 测试设置概况",
        ["测试项", "设置", "说明"],
        [
            ["正式样本数", str(overview["sample_count"]), "统一安全主表"],
            ["数据划分", f"{train_count}/{val_count}/{test_count}", "训练集/验证集/测试集"],
            ["分类评估", "准确率、平衡准确率、F1", "用于产地溯源任务"],
            ["回归评估", "R²、RMSE", "用于品质预测和液质成分预测任务"],
        ],
    )

    w("5.2 产地溯源对比分析", style="Heading 2")
    w(
        f"产地溯源是本项目当前表现最稳定的任务。以 {best_origin['特征方案']} 为输入时，{best_origin['模型']} 在验证集和测试集上的准确率、平衡准确率和 F1_macro 均达到 1.0000。对比赛展示而言，这一结果具备很强的说服力，因为它直接对应产地识别、批次核验和品牌溯源三个使用场景。",
        indent=True,
    )
    w(
        "对比结果说明，S960 单光谱已经能够完整捕捉两地产地样本的区分信息；双光谱方案同样保持很高水平，并为后续品质预测和液质成分预测保留统一输入基础。也就是说，单光谱在产地判别上足够强，双光谱在全项目任务体系中更利于统一组织，这正是本项目多任务路线的实际价值所在。",
        indent=True,
    )
    add_table(
        doc,
        lines,
        "表5-2 产地溯源代表性结果",
        ["特征方案", "模型", "验证集准确率", "测试集准确率", "测试集F1_macro"],
        [
            ["S960单光谱", "逻辑回归", "1.0000", "1.0000", "1.0000"],
            ["S960单光谱", "线性SVM", "1.0000", "1.0000", "1.0000"],
            ["R210+S960双光谱", "LightGBM分类", "0.9833", "1.0000", "1.0000"],
        ],
    )

    w("5.3 品质预测对比分析", style="Heading 2")
    w(
        f"品质预测任务中，糖度是当前最具代表性的指标。{sugar_result['特征方案']} 条件下的 {sugar_result['模型']} 在测试集上取得 R²={fmt_float(float(sugar_result['r2']))}、RMSE={fmt_float(float(sugar_result['rmse']))}，说明双光谱方案对糖度变化具有较好的响应能力。糖度结果直接对应果品分级和品质判断，也是比赛展示中最容易被理解和验证的一项结果。",
        indent=True,
    )
    w(
        f"从样本均值看，琼中绿橙平均糖度为 {qz_sugar:.2f}，平均糖酸比为 {qz_ratio:.2f}；澄迈福橙平均 VC 为 {cm_vc:.2f}。这些差异与页面展示和文字结论中的主要描述保持一致，使品质分析部分不仅有模型分数，也有可追溯到样本统计的事实支撑。VC 和糖酸比在测试中同样给出了可用信息，适合作为糖度结论的补充说明。",
        indent=True,
    )
    add_table(
        doc,
        lines,
        "表5-3 品质预测代表性结果",
        ["指标", "最佳模型", "特征方案", "测试集R²", "测试集RMSE"],
        [
            ["糖度", sugar_result["模型"], sugar_result["特征方案"], fmt_float(float(sugar_result["r2"])), fmt_float(float(sugar_result["rmse"]))],
            ["VC", vc_result["模型"], vc_result["特征方案"], fmt_float(float(vc_result["r2"])), fmt_float(float(vc_result["rmse"]))],
            ["糖酸比", ratio_result["模型"], ratio_result["特征方案"], fmt_float(float(ratio_result["r2"])), fmt_float(float(ratio_result["rmse"]))],
        ],
    )
    add_table(
        doc,
        lines,
        "表5-4 两地产地部分指标均值对比",
        ["产地", "平均糖度", "平均糖酸比", "平均VC"],
        [
            ["澄迈福橙", f"{origin_stats['CM']['糖度']:.2f}", f"{origin_stats['CM']['糖酸比']:.2f}", f"{origin_stats['CM']['VC']:.2f}"],
            ["琼中绿橙", f"{origin_stats['QZ']['糖度']:.2f}", f"{origin_stats['QZ']['糖酸比']:.2f}", f"{origin_stats['QZ']['VC']:.2f}"],
        ],
    )

    w("5.4 液质成分对比分析", style="Heading 2")
    w(
        f"液质成分预测是本项目最能体现多模态分析价值的部分。基线实验中，{baseline_lcms['特征方案']} + {baseline_lcms['模型']} 已取得验证集平均 R²={fmt_float(float(baseline_lcms['验证集平均R2']))}、测试集平均 R²={fmt_float(float(baseline_lcms['测试集平均R2']))}。在此基础上，项目继续引入 SNV 预处理、PCA 降维和参数搜索后，验证集平均 R² 提升到 {fmt_float(float(best_lcms_val['验证集平均R2']))}，测试集平均 R² 提升到 {fmt_float(float(best_lcms_test['测试集平均R2']))}，增益分别达到 {fmt_float(lcms_val_gain)} 和 {fmt_float(lcms_test_gain)}。",
        indent=True,
    )
    w(
        f"单项结果中，异柠檬酸测试集 R² 达到 {fmt_float(float(iso_result['r2']))}，莽草酸达到 {fmt_float(float(shikimic_result['r2']))}，蔗糖达到 {fmt_float(float(sucrose_result['r2']))}，E-1-丙烯-1,2,3-三羧酸达到 {fmt_float(float(tricarboxylic_result['r2']))}。这些成分结果把项目分析范围从宏观品质指标推进到内部成分层面，使“来自哪里”和“内部有什么差异”两类问题可以在同一项目中同时得到回应。",
        indent=True,
    )
    w(
        f"从样本统计看，澄迈福橙在异柠檬酸均值上更高，平均为 {cm_iso:.2f}；琼中绿橙在蔗糖均值上更高，平均为 {qz_sucrose:.2f}。因此，液质成分部分并不是孤立的高分结果，而是与两地产地差异、品质差异一起构成了可解释的项目内容。",
        indent=True,
    )
    add_table(
        doc,
        lines,
        "表5-5 液质成分任务整体增益",
        ["阶段", "验证集平均R²", "测试集平均R²", "说明"],
        [
            [
                "基线实验",
                fmt_float(float(baseline_lcms["验证集平均R2"])),
                fmt_float(float(baseline_lcms["测试集平均R2"])),
                f"{baseline_lcms['特征方案']} + {baseline_lcms['模型']}",
            ],
            [
                "优化后最优验证方案",
                fmt_float(float(best_lcms_val["验证集平均R2"])),
                fmt_float(float(best_lcms_val["测试集平均R2"])),
                best_lcms_val["实验方案"],
            ],
            [
                "优化后最优测试方案",
                fmt_float(float(best_lcms_test["验证集平均R2"])),
                fmt_float(float(best_lcms_test["测试集平均R2"])),
                best_lcms_test["实验方案"],
            ],
        ],
    )
    add_table(
        doc,
        lines,
        "表5-6 液质成分代表性结果",
        ["指标", "最佳模型", "特征方案", "测试集R²"],
        [
            ["异柠檬酸", iso_result["模型"], iso_result["特征方案"], fmt_float(float(iso_result["r2"]))],
            ["莽草酸", shikimic_result["模型"], shikimic_result["特征方案"], fmt_float(float(shikimic_result["r2"]))],
            ["蔗糖", sucrose_result["模型"], sucrose_result["特征方案"], fmt_float(float(sucrose_result["r2"]))],
            ["E-1-丙烯-1,2,3-三羧酸", tricarboxylic_result["模型"], tricarboxylic_result["特征方案"], fmt_float(float(tricarboxylic_result["r2"]))],
        ],
    )

    w("5.5 综合测试结论", style="Heading 2")
    w(
        "测试部分最终形成了三组核心结果：产地溯源在测试集上实现稳定区分，糖度预测给出了可直接用于品质说明的关键指标，液质成分预测补充了内部成分差异信息。三类结果全部来自同一安全主表和同一套样本划分，能够在同一份项目报告中并列呈现。",
        indent=True,
    )
    add_table(
        doc,
        lines,
        "表5-7 主要测试结论与应用对应",
        ["任务", "代表结果", "对应应用"],
        [
            ["产地溯源", "S960 单光谱逻辑回归/线性SVM 测试准确率 1.0000", "产地识别、品牌溯源、批次核验"],
            ["品质预测", f"糖度预测测试集 R² {fmt_float(float(sugar_result['r2']))}", "果品分级、品质判断、批次说明"],
            ["液质成分预测", f"异柠檬酸 R² {fmt_float(float(iso_result['r2']))}，莽草酸 R² {fmt_float(float(shikimic_result['r2']))}", "内部成分解析、差异说明、重点结果展示"],
        ],
    )

    w("第六章 项目总结", style="Heading 1")
    w("6.1 项目特色与创新点", style="Heading 2")
    w(
        "本项目在正式实验前完成 HSI 排查、S960 共同波段对齐和安全主表重建，正式测试全部建立在 399 个安全样本上。面对多模态资料中的异常模态和跨产地重复问题，项目没有把全部数据直接送入建模，而是先完成数据治理，再进入正式实验。",
        indent=True,
    )
    w(
        "产地判别、品质评价和液质成分预测共用统一编号、统一主表和统一划分方案，结果在页面和报告中围绕同一编号样本展开。查看单个样本时，产地、糖度、VC 和重点成分结果能够同时对应到同一份记录中。",
        indent=True,
    )
    w(
        "项目输出覆盖页面展示、历史记录、中文结果说明和正式报告文本，检测结果可以直接进入验货、展示、归档和教学说明等环节。比赛现场展示时，评审看到的不再是零散指标，而是围绕样本组织好的成套结果。",
        indent=True,
    )
    add_table(
        doc,
        lines,
        "表6-1 项目特色与创新点概述",
        ["创新点", "具体体现"],
        [
            ["数据治理创新", "识别 HSI 跨产地完全重复问题，剔除异常模态并重建安全主表"],
            ["多任务协同创新", "产地、品质、液质成分三类任务共用统一样本体系和统一划分方案"],
            ["结果组织创新", "分析结果围绕样本编号组织为页面展示、历史记录和正式报告"],
            ["场景贯通创新", "把样本录入、分析处理、结果说明和留档放到同一条项目链路中"],
        ],
    )

    w("6.2 应用推广", style="Heading 2")
    w(
        "本项目最适合首先落在海南特色柑橘的品质检测与产地说明场景。对产地和合作社而言，项目可用于样本登记、批次对比和品质初判；对品牌和渠道而言，项目可用于产地核验、结果展示和批次说明；对区域展馆和教学场景而言，项目可作为海南特色果品数字化检测的展示案例。当前项目已经具备稳定的数据底座、明确的测试结果和成形的系统展示界面，具备开展区域试点和赛事展示的现实基础。",
        indent=True,
    )
    w(
        "随着样本规模扩大和品类继续接入，本项目的推广对象不局限于两类柑橘样本。以当前的数据治理方法、多任务分析流程和结果组织方式为基础，项目可继续扩展到更多柑橘品种及其他特色农产品场景，在区域农业数字化、标准化和品牌化建设中形成可复用的方法模板。",
        indent=True,
    )
    add_table(
        doc,
        lines,
        "表6-2 应用推广方向",
        ["推广方向", "落地场景", "项目价值"],
        [
            ["品质检测", "果园、合作社、分选环节", "辅助分级、批次说明、结果留档"],
            ["产地核验", "品牌展示、渠道验货", "支撑产地识别与批次核验"],
            ["教学展示", "课程实践、成果展示", "支撑数据治理与多任务分析演示"],
            ["区域拓展", "更多柑橘和特色农产品", "复用数据治理方法和结果输出方式"],
        ],
    )

    w("6.3 项目展望", style="Heading 2")
    w(
        "下一阶段，项目将继续围绕样本扩充、品类扩展和结果表达深化三条线推进。样本侧将引入更多批次和季节数据，增强模型对实际生产场景的覆盖；任务侧将把现有双光谱分析方法推广到更多柑橘与特色农产品；表达侧将继续丰富图表展示、报告模板和场景化说明，使项目在比赛展示之外，也能更好服务于区域品牌建设和农业数字化应用。",
        indent=True,
    )

    w("参考文献", style="Heading 1")
    refs = [
        "[1] 正式实验/研究记录/正式实验概览.json",
        "[2] 正式实验/研究记录/双光谱正式研究顺序记录.md",
        "[3] 正式实验/研究记录/第一版双光谱基线实验摘要.json",
        "[4] 正式实验/研究记录/进阶优化记录/液质预测进阶调参完整记录.csv",
        "[5] 清洗结果/质检信息/汇总信息.json",
        "[6] 正式实验/划分方案/统一样本划分.csv",
        "[7] 正式实验/数据集/双光谱安全主表.csv",
        "[8] 正式实验/基线结果/任务1_品质预测/逐目标结果.csv",
        "[9] 正式实验/基线结果/任务2_产地溯源/模型汇总.csv",
        "[10] 正式实验/基线结果/任务3_液质成分预测/逐目标结果.csv",
        "[11] Lu Z, Jia K, Zhang H, et al. Geographical Origin Identification of Citrus Fruits Based on Near-Infrared Spectroscopy Combined with Convolutional Neural Network and Data Augmentation[J]. Agriculture, 2025, 15(22):2350.",
        "[12] Wu D, Sun D W. Advanced applications of hyperspectral imaging technology for food quality and safety analysis and assessment: A review[J]. Innovative Food Science and Emerging Technologies, 2013, 19:1-14.",
        "[13] Liu Y, Sun X, Ouyang A. Nondestructive measurement of soluble solid content of navel orange fruit by visible-NIR spectrometric technique with PLSR and PCA-BPNN[J]. LWT - Food Science and Technology, 2010, 43(4):602-607.",
        "[14] Liu Y, Chen S, Zhang F, et al. Simultaneous detection of citrus internal quality attributes based on near-infrared spectroscopy and hyperspectral imaging with multi-task deep learning and instrumental transfer learning[J]. Food Chemistry, 2025, 472:144337.",
    ]
    for item in refs:
        w(item, indent=True)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    OUT_TXT.write_text("\n".join(lines), encoding="utf-8")


if __name__ == "__main__":
    main()

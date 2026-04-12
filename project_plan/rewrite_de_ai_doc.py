from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


SRC = Path(r"D:\Projects\chengyuan-web\project_plan\_tmp_v33_utf8.txt")
OUT_TXT = Path(r"D:\Projects\chengyuan-web\project_plan\chengyuan_de_ai_rewrite_20260410.txt")
OUT_DOCX = Path(r"D:\Projects\chengyuan-web\project_plan\chengyuan_de_ai_rewrite_20260410.docx")


def is_number_heading(text: str) -> bool:
    if not text or not text[0].isdigit() or "." not in text:
        return False
    head = text.split(" ", 1)[0]
    parts = [p for p in head.split(".") if p]
    return bool(parts) and all(p.isdigit() for p in parts)


def main() -> None:
    lines = SRC.read_text(encoding="utf-8").splitlines()

    lines[0] = "《橙源智鉴——基于多模态光谱融合与深度学习的海南柑橘品质智能检测与产地溯源平台》"
    lines[1] = "文本修订版"
    lines[2] = "修订时间：2026年4月10日"

    summary = (
        "我国柑橘年产量超6000万吨，品牌果、电商果、礼盒果和渠道采购持续增长，"
        "但产业端仍普遍存在等级判定靠经验、产地证明停留在纸面、检测结果难真正进入"
        "分选验货和品牌经营等问题。围绕这一行业痛点，本项目打造“橙源智鉴”智能分析平台，"
        "面向橙类产品构建数据文件上传、等级识别、产地溯源、品质分析、异常提示、历史查询和"
        "报告输出一体化解决方案。项目已完成网站端、小程序端与标准报告体系联动，可支持 HDR、"
        "RAW、CSV 等多类文件接入，形成“样本采集-理化标定-模型训练-双端分析-报告归档”的"
        "完整闭环。现阶段已沉淀 10 组化验样本、100 个连续波段点位及澄迈福橙、琼中绿橙等"
        "区域对比数据；在演示场景中，单次分析流程约 1.2 秒，产地识别置信度可达 98.5%，"
        "SSC 预测误差展示为 ±0.3%。与传统只输出一次结论的检测方案不同，橙源智鉴把批次结果、"
        "历史档案、图文化报告和场景交付打通，既能服务果园、合作社和分选中心的分级验货，"
        "也能服务品牌方和渠道方的品质说明、溯源展示与复检复盘。项目的重点不在于堆叠算法概念，"
        "而在于把现有样本、分析结果和交付流程做成可展示、可试点、可继续验证的数字化品控工具。"
    )
    lines[66] = "摘要"
    lines.insert(67, summary)

    lines[113] = "　　定性识别路径（测等级）：采用多分类算法快速判定样品等级与产地。在当前演示样本中，产地识别置信度可达98.5%。"
    lines[114] = "　　定量预测路径（测指标）：采用回归算法建立特征光谱与糖度、酸度等指标的非线性映射关系，当前页面展示的 SSC 预测误差为 ±0.3%。"

    lines[182] = "　　市场可行性的重点，不在于把市场写得多大，而在于需求是否真实、客户是否明确、项目是否已经接近实际成交场景。橙类品质检测、批次分级和产地表达都不是被创造出来的新概念，而是果园、合作社、分选中心、品牌方和渠道采购方每天都会发生的现实动作。橙源智鉴把这些原本分散在线下经验、纸面记录和零散检测中的动作整合成可直接调用的平台流程，因此它面对的是明确存在且高频复发的经营需求。"
    lines[285] = "　　技术章节的重点，不在于堆叠算法名词，而在于说明技术如何进入产品并形成交付。橙源智鉴的技术路线不是单一模型展示，而是围绕高光谱文件进入系统后的完整处理链搭建：前端完成数据上传和批次关联，后端完成预处理、波段筛选、多任务建模和异常识别，结果端同步输出等级结论、产地判断、品质指标和标准报告，使模型结果能够直接服务分级、验货和品牌说明等业务动作。"
    lines[306] = "　　产品策略部分需要避免把系统写成若干功能点的简单堆砌。橙源智鉴当前已经形成可交付产品形态，交付的不只是一个网站页面或一个小程序入口，而是“上传分析平台 + 移动查询端 + 批次报告模板 + 历史档案体系”四位一体的成果包。合作方拿到的不只是演示结果，而是一套能进入试点、能被现场人员使用、能被品牌方对外调用的完整工具。"
    lines[340] = "　　团队介绍部分更需要说明“这支队伍为什么能把事情做成”，而不是只把组织名称列得完整。橙源智鉴团队的优势，在于它不是单一技术团队，而是把样本组织、理化标定、模型训练、平台开发、场景对接和报告交付串成了一支能完成闭环交付的复合型队伍。团队成员来自人工智能、软件工程、计算机科学与技术等相关专业，既能处理高光谱建模与系统开发，也能承接试点沟通、材料输出和项目推进，具备较强的工程落地能力。"
    lines[391] = "　　彭斯逸同学性格开朗、做事认真，具备较强的组织协调与团队协作能力，积极组织并参与各项班级、学院活动。自入学以来学业表现稳定，专业排名前3%，核心课程均分90+，曾获中国国际大学生创新大赛校级金奖、全国大学生三创赛校级一等奖等荣誉。"
    lines[392] = "　　在技术方向上，他持续参与AI应用开发与后端工程实践，熟悉Python异步编程及FastAPI框架，在检索增强生成（RAG）方向有相关积累，曾受邀担任Datawhale“RAG技术全栈指南”课程专业助教。他曾于软件公司研发部门实习，参与垂直领域大模型应用落地及数据清洗管线设计，积累了一定的工程实践经验。"

    for idx, line in enumerate(lines):
        if "综合毛利率达到60%以上" in line:
            lines[idx] = "　　这种“轻硬件、重算法”的产品结构，意味着项目在服务规模扩大后，单位交付成本有望逐步下降，并为后续渠道拓展、售后服务及研发迭代留出空间。"
        if "可在设备投入使用后的数月内全面收回初始采购成本" in line:
            lines[idx] = "　　从使用逻辑看，终端用户若能稳定把系统用于分级、验货和复检，前期投入有机会在较短周期内逐步摊薄。具体回收节奏仍取决于客户体量、使用频次和交付方式。"

    text = "\n".join(lines)
    text = re.sub(r"\n{3,}", "\n\n", text).strip() + "\n"
    OUT_TXT.write_text(text, encoding="utf-8")

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "宋体"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    for i, line in enumerate(text.splitlines()):
        raw = line.strip()
        if not raw:
            doc.add_paragraph("")
            continue
        if i == 0:
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(raw)
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            continue
        if raw in {"文本修订版", "修订时间：2026年4月10日"}:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(raw)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(11)
            continue
        if raw in {"目录", "摘要"} or (raw.startswith("第") and "章" in raw[:5]):
            doc.add_paragraph(raw, style="Heading 1")
            continue
        if is_number_heading(raw):
            style = "Heading 3" if raw.split(" ", 1)[0].count(".") >= 2 else "Heading 2"
            doc.add_paragraph(raw, style=style)
            continue
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(24)
        run = p.add_run(raw)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(12)

    doc.save(OUT_DOCX)
    print(OUT_TXT)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
